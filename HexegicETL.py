import pandas as pd
from sodapy import Socrata
import docx
import os
import shutil
import zipfile


def extract():
    """
    As suggested in the API Docs https://dev.socrata.com/foundry/healthdata.gov/j8mb-icvb
    :return: The complete dataframe from https://healthdata.gov/dataset/COVID-19-Diagnostic-Laboratory-Testing-PCR-Testing/j8mb-icvb
    """
    client = Socrata('healthdata.gov', None)
    # results = client.get_all('j8mb-icvb') # Safer option for picking up all the data, however significantly slows down the process
    results = client.get('j8mb-icvb', limit=1000000000)
    df = pd.DataFrame.from_records(results)

    return df


def transform(df):
    """
    Performing the desired dataframe transformations
    :param df: Dataframe with the following columns 'state', 'state_name', 'state_fips', 'fema_region', 'overall_outcome',
                'date', 'new_results_reported', 'total_results_reported'
    :return: Transformed Dataframe with the following columns 'Date', 'state_name', 'New Positive Tests Reported',
                'Total Positive Tests Reported', 'New Negative Tests Reported', 'Total Negative Tests Reported'
    """
    positive_df = df[df['overall_outcome'] == 'Positive']
    positive_df = positive_df[['date', 'state_name', 'new_results_reported', 'total_results_reported']]
    positive_df.rename(columns={'date': 'Date', 'new_results_reported': 'New Positive Tests Reported',
                                'total_results_reported': 'Total Positive Tests Reported'},
                       inplace=True)

    negative_df = df[df['overall_outcome'] == 'Negative']
    negative_df = negative_df[['date', 'state_name', 'new_results_reported', 'total_results_reported']]
    negative_df.rename(columns={'date': 'Date', 'new_results_reported': 'New Negative Tests Reported',
                                'total_results_reported': 'Total Negative Tests Reported'},
                       inplace=True)

    transformed_df = pd.merge(positive_df, negative_df, on=['Date', 'state_name'], how='outer')
    transformed_df['Date'] = pd.to_datetime(transformed_df.Date).astype('str')

    return transformed_df


def load(df):
    """
    Construct a zip file with all the desired documents (.docx) for each State
    :param df: Dataframe in final form, with the addition of the column state_name (which will be used to filter)
    :return: Zip File in the scripts directory
    """
    state_names = ['Alabama', 'Alaska', 'Arizona', 'Arkansas', 'California', 'Colorado', 'Connecticut', 'Delaware',
                   'Florida', 'Georgia', 'Hawaii', 'Idaho', 'Illinois', 'Indiana', 'Iowa', 'Kansas', 'Kentucky',
                   'Louisiana', 'Maine', 'Maryland', 'Massachusetts', 'Michigan', 'Minnesota', 'Mississippi',
                   'Missouri', 'Montana', 'Nebraska', 'Nevada', 'New Hampshire', 'New Jersey', 'New Mexico', 'New York',
                   'North Carolina', 'North Dakota', 'Ohio', 'Oklahoma', 'Oregon', 'Pennsylvania', 'Rhode Island',
                   'South Carolina', 'South Dakota', 'Tennessee', 'Texas', 'Utah', 'Vermont', 'Virginia', 'Washington',
                   'West Virginia', 'Wisconsin', 'Wyoming']

    run_date = pd.Timestamp.today().strftime('%Y-%m-%d')
    directory = f'{os.getcwd()}/US_COVID-19_Testing_{run_date}'
    if not os.path.exists(directory):
        os.mkdir(directory)

    for state in state_names:
        doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(12)

        doc.add_heading('The current status of COVID-19 testing', 1)
        doc.add_heading(state, 0)

        p = doc.add_paragraph()
        p.style = doc.styles.add_style('Style Name', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        font = p.style.font
        font.name = 'Arial'
        font.size = docx.shared.Pt(9)
        p.add_run('Date of the Report: ').bold = True
        p.add_run(run_date)

        doc.add_paragraph('')

        doc.add_heading('Data for the 5 most recent days’ worth of lab test results available', 4)

        state_df = df[(df['state_name'] == state)] \
            .drop(['state_name'], axis=1) \
            .sort_values('Date')\
            .tail(5)

        t = doc.add_table(rows=1, cols=state_df.shape[1])
        for j in range(state_df.shape[1]):
            cell = state_df.columns[j]
            p = t.cell(0, j).add_paragraph('')
            p.add_run(str(cell)).bold = True

        for i in range(state_df.shape[0]):
            row = t.add_row()
            for j in range(state_df.shape[1]):
                cell = state_df.iat[i, j]
                row.cells[j].text = str(cell)

        doc.save(f'{directory}/{state}_{run_date}.docx')

    with zipfile.ZipFile(f'US_COVID-19_Testing_{run_date}.zip', 'w') as zipf:
        for f in os.listdir(directory + '/'):
            zipf.write(f'{directory}/{f}', f)
        zipf.close()

    shutil.rmtree(directory)


if __name__ == "__main__":
    extract_data_df = extract()
    transformed_data_df = transform(extract_data_df)
    load(transformed_data_df)
